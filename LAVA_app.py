from __future__ import annotations

import io
import streamlit as st
import requests
import pandas as pd
from io import BytesIO

# ── Secrets ───────────────────────────────────────────────────────────────────
TOKEN = st.secrets["GITHUB_PAT"]
REPO  = f"{st.secrets['LAVA_OWNER']}/{st.secrets['LAVA_REPO']}"
HEADERS = {
    "Authorization": f"token {TOKEN}",
    "Accept": "application/vnd.github.v3+json",
}

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="LAVA — Lecture Attendance Archive",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.hero {
    background: linear-gradient(135deg, #7b0000 0%, #c0392b 100%);
    border-radius: 16px; padding: 2rem; text-align: center;
    color: white; margin-bottom: 2rem;
    box-shadow: 0 4px 20px rgba(123,0,0,0.3);
}
.hero h1 { font-size: 2rem; font-weight: 900; margin: 0; }
.hero .sub { opacity: 0.85; margin-top: 0.3rem; font-size: 0.95rem; }
.hero .badge {
    display: inline-block; background: rgba(255,255,255,0.2);
    border-radius: 20px; padding: 3px 14px; font-size: 0.8rem; margin-top: 0.6rem;
}
.info-card {
    background: rgba(192, 57, 43, 0.08);
    border-left: 4px solid #c0392b;
    border-radius: 0 8px 8px 0;
    padding: 0.7rem 1rem; margin: 0.5rem 0;
    font-size: 0.9rem; color: inherit;
}
.info-card b { color: #c0392b; }
.stat-box {
    background: rgba(192, 57, 43, 0.07);
    border: 1px solid rgba(192, 57, 43, 0.2);
    border-radius: 10px; padding: 1rem;
    text-align: center; color: inherit;
}
.stat-box .num { font-size: 2rem; font-weight: 900; color: #c0392b; }
.stat-box .lbl { font-size: 0.8rem; opacity: 0.7; text-transform: uppercase; letter-spacing: 1px; }
div[data-testid="stForm"] {
    border: 1.5px solid rgba(128,128,128,0.3);
    border-radius: 12px; padding: 1.2rem 1.2rem 0.5rem;
}
</style>
""", unsafe_allow_html=True)

# ── Fixed footer ─────────────────────────────────────────────────────────────
st.markdown('\n<style>\n.ulas-footer {\n    position: fixed;\n    bottom: 0; left: 0; right: 0;\n    text-align: center;\n    padding: 0.45rem 1rem;\n    font-size: 0.78rem;\n    background: rgba(0,0,0,0.45);\n    backdrop-filter: blur(6px);\n    -webkit-backdrop-filter: blur(6px);\n    color: rgba(255,255,255,0.55);\n    letter-spacing: 0.04em;\n    z-index: 9999;\n    border-top: 1px solid rgba(255,255,255,0.07);\n}\n.ulas-footer b { color: rgba(255,255,255,0.8); font-weight: 600; }\n.ulas-footer .dot { color: rgba(255,255,255,0.3); margin: 0 0.3em; }\n/* Push content up so footer never overlaps last element */\nsection.main > div { padding-bottom: 2.8rem !important; }\n</style>\n<div class="ulas-footer">\n    Made with ❤️ by\n    <b>SESET</b><span class="dot">•</span><b>EPE</b><span class="dot">•</span><b>2030/2031</b>\n</div>\n', unsafe_allow_html=True)

st.markdown("""
<div class="hero">
    <h1>📚 LAVA</h1>
    <div class="sub">Lecture Attendance Viewing Archive</div>
    <div class="badge">Federal University of Technology, Owerri</div>
</div>
""", unsafe_allow_html=True)


# ── GitHub helpers ─────────────────────────────────────────────────────────────
@st.cache_data(ttl=60, show_spinner=False)
def gh_list(path: str) -> list[dict]:
    url = f"https://api.github.com/repos/{REPO}/contents/{path}"
    r   = requests.get(url, headers=HEADERS, timeout=10)
    if r.status_code != 200:
        return []
    data = r.json()
    return data if isinstance(data, list) else []

@st.cache_data(ttl=300, show_spinner=False)
def fetch_csv_bytes(download_url: str) -> bytes:
    return requests.get(download_url, headers=HEADERS, timeout=10).content

def list_dirs(path: str) -> list[str]:
    return sorted(
        [f["name"] for f in gh_list(path) if f.get("type") == "dir"],
        reverse=True,
    )

def list_csvs(path: str) -> list[dict]:
    return [f for f in gh_list(path) if f.get("name", "").endswith(".csv")]


# ── Export helpers ─────────────────────────────────────────────────────────────
def df_to_excel(df: pd.DataFrame, title: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Attendance"

    # Title
    ws.merge_cells("A1:E1")
    ws["A1"] = title
    ws["A1"].font      = Font(name="Arial", bold=True, size=13)
    ws["A1"].alignment = Alignment(horizontal="center")

    # Blank spacer row 2
    headers    = list(df.columns)
    thin       = Side(style="thin", color="CCCCCC")
    border     = Border(left=thin, right=thin, top=thin, bottom=thin)
    hdr_fill   = PatternFill("solid", fgColor="C0392B")
    alt_fill   = PatternFill("solid", fgColor="F9EBEA")

    # Header row 3
    for col_i, h in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=col_i, value=h)
        cell.font      = Font(name="Arial", bold=True, color="FFFFFF", size=10)
        cell.fill      = hdr_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border    = border

    # Data rows starting at row 4
    for row_i, row in enumerate(df.itertuples(index=False), start=4):
        for col_i, value in enumerate(row, start=1):
            cell = ws.cell(row=row_i, column=col_i, value=value)
            cell.font      = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="left")
            cell.border    = border
            if row_i % 2 == 0:
                cell.fill = alt_fill

    # Column widths
    widths = {"S/N": 6, "Surname": 22, "Other Names": 26, "Matric Number": 16, "Time": 20}
    for i, col in enumerate(headers, start=1):
        letter = ws.cell(row=3, column=i).column_letter
        ws.column_dimensions[letter].width = widths.get(col, 18)

    # Total row
    last_data = len(df) + 3
    total_row = last_data + 1
    tc = ws.cell(row=total_row, column=1, value="Total Students")
    tc.font = Font(name="Arial", bold=True, size=10)
    tv = ws.cell(row=total_row, column=2, value=f"=COUNTA(B4:B{last_data})")
    tv.font = Font(name="Arial", bold=True, size=10)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def df_to_word(df: pd.DataFrame, title: str, filename: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"),  hex_color)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"),   "clear")
        tcPr.append(shd)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Institution header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FEDERAL UNIVERSITY OF TECHNOLOGY, OWERRI")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x7B, 0x00, 0x00)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Lecture Attendance Record")
    r2.italic = True
    r2.font.size = Pt(11)

    doc.add_paragraph()

    # Parse filename: SCHOOLABBRLEVEL_COURSECODE_DATE.csv
    try:
        base   = filename.replace(".csv", "")
        parts  = base.split("_")
        ident  = parts[0]
        course = parts[1] if len(parts) > 1 else "—"
        date   = parts[2] if len(parts) > 2 else "—"
    except Exception:
        ident = course = date = "—"

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(f"{course}  |  {ident}  |  {date}")
    rt.bold = True
    rt.font.size = Pt(12)

    doc.add_paragraph()

    # Table
    headers = list(df.columns)
    table   = doc.add_table(rows=1, cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header cells
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "C0392B")

    # Data rows
    for row_i, row in enumerate(df.itertuples(index=False)):
        cells = table.add_row().cells
        for col_i, value in enumerate(row):
            cells[col_i].text = str(value)
            cells[col_i].paragraphs[0].runs[0].font.size = Pt(9)
            if row_i % 2 == 0:
                set_cell_bg(cells[col_i], "F9EBEA")

    # Column widths
    col_widths = [Cm(1.2), Cm(4.5), Cm(5.0), Cm(3.5), Cm(3.5)]
    for i, col in enumerate(table.columns):
        if i < len(col_widths):
            for cell in col.cells:
                cell.width = col_widths[i]

    doc.add_paragraph()
    fp = doc.add_paragraph(f"Total Students: {len(df)}")
    fp.runs[0].bold = True
    fp.runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Browse ─────────────────────────────────────────────────────────────────────
ROOT  = "attendances"
dates = list_dirs(ROOT)

if not dates:
    st.info("No attendance records have been pushed to LAVA yet.")
    st.caption("Records appear automatically when course reps end attendance sessions in ULAS.")
    st.stop()

st.markdown("### 📅 Select Date")
selected_date = st.selectbox("Date", dates)

with st.spinner(f"Loading records for {selected_date}..."):
    all_files = list_csvs(f"{ROOT}/{selected_date}")

if not all_files:
    st.warning(f"No attendance records found for {selected_date}.")
    st.stop()

# ── Search ─────────────────────────────────────────────────────────────────────
st.markdown("### 🔍 Filter")
search = st.text_input(
    "Search by course code, abbreviation, school, or level",
    placeholder="e.g. CSC301  or  EEE300  or  SEET",
)

filtered = all_files
if search.strip():
    filtered = [f for f in all_files if search.strip().lower() in f["name"].lower()]

if not filtered:
    st.warning("No records match your search.")
    st.stop()

# ── Stats ──────────────────────────────────────────────────────────────────────
st.markdown("<br>", unsafe_allow_html=True)
s1, s2 = st.columns(2)
with s1:
    st.markdown(
        f'<div class="stat-box"><div class="num">{len(filtered)}</div>'
        f'<div class="lbl">Records for {selected_date}</div></div>',
        unsafe_allow_html=True,
    )
with s2:
    courses = set()
    for f in filtered:
        parts = f["name"].replace(".csv", "").split("_")
        if len(parts) >= 2:
            courses.add(parts[1])
    st.markdown(
        f'<div class="stat-box"><div class="num">{len(courses)}</div>'
        f'<div class="lbl">Courses</div></div>',
        unsafe_allow_html=True,
    )

st.markdown("<br>", unsafe_allow_html=True)

# ── File picker ────────────────────────────────────────────────────────────────
# Filename: SCHOOLABBRLEVEL_COURSECODE_DATE.csv
def make_label(name: str) -> str:
    try:
        base  = name.replace(".csv", "")
        parts = base.split("_")
        return f"{parts[2]}  ·  {parts[1]}  ·  {parts[0]}"
    except Exception:
        return name

st.markdown("### 📄 Select Attendance Record")
file_map       = {make_label(f["name"]): f for f in sorted(filtered, key=lambda x: x["name"])}
selected_label = st.selectbox("Attendance file", list(file_map.keys()))
selected_file  = file_map[selected_label]
filename       = selected_file["name"]

# Parse for display
try:
    base   = filename.replace(".csv", "")
    parts  = base.split("_")
    ident  = parts[0]
    course = parts[1] if len(parts) > 1 else "—"
    date   = parts[2] if len(parts) > 2 else "—"
except Exception:
    ident = course = date = "—"

st.markdown(f"""<div class="info-card">
    <b>Reference:</b> {ident} &nbsp;|&nbsp;
    <b>Course:</b> {course} &nbsp;|&nbsp;
    <b>Date:</b> {date}
</div>""", unsafe_allow_html=True)

# ── Load & display ─────────────────────────────────────────────────────────────
with st.spinner("Loading..."):
    csv_bytes = fetch_csv_bytes(selected_file["download_url"])

try:
    df = pd.read_csv(BytesIO(csv_bytes))
except Exception as e:
    st.error(f"Could not read CSV: {e}")
    st.stop()

doc_title = f"Attendance — {course} | {ident} | {date}"
base_name = filename.replace(".csv", "")

st.markdown(f"### {course} — {ident} &nbsp;&nbsp; ({len(df)} students)")
st.dataframe(df, use_container_width=True, hide_index=True)

# ── Download buttons ───────────────────────────────────────────────────────────
st.markdown("#### Download Record")
dl1, dl2, dl3 = st.columns(3)

with dl1:
    st.download_button(
        "📄 CSV",
        csv_bytes,
        file_name=filename,
        mime="text/csv",
        use_container_width=True,
    )

with dl2:
    try:
        xlsx_bytes = df_to_excel(df, doc_title)
        st.download_button(
            "📊 Excel",
            xlsx_bytes,
            file_name=f"{base_name}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    except Exception as e:
        st.error(f"Excel error: {e}")

with dl3:
    try:
        docx_bytes = df_to_word(df, doc_title, filename)
        st.download_button(
            "📝 Word",
            docx_bytes,
            file_name=f"{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except Exception as e:
        st.error(f"Word error: {e}")

# ── Footer ─────────────────────────────────────────────────────────────────────

